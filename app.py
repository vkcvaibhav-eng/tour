import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import os
import json
import pandas as pd
from datetime import datetime
import io

# --- CONFIGURE PAGE ---
st.set_page_config(page_title="NAU Tour Diary Generator", layout="wide")

st.title("📝 Automated Tour Diary Generator (Landscape)")
st.markdown("""
**Generates a NAU Tour Diary in Landscape format matching the official Letter style.**
* **Upload:** Tour Orders (contains System No.), Google Maps Screenshots (for Distance), and Tickets.
* **Smart Merge:** Automatically extracts distance from Map screenshots and System IDs from Tour Orders.
* **Output:** Formatted .docx with required Header, B.H., and Signature blocks.
""")

# --- SIDEBAR: API KEY ---
with st.sidebar:
    st.header("🔑 API Configuration")
    GEMINI_API_KEY = st.text_input("Gemini API Key", type="password")
    st.info("Get your key from Google AI Studio.")

# --- HELPER FUNCTIONS ---

def set_landscape(doc):
    """Sets the document section to landscape orientation."""
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # Adjust margins for landscape
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

def extract_doc_data(uploaded_file, api_key):
    """
    Uses Gemini to extract data from Tour Orders, Tickets, Salary Slips, or Map Screenshots.
    """
    genai.configure(api_key=api_key)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.getvalue())
        tmp_path = tmp.name

    try:
        sample_file = genai.upload_file(path=tmp_path, display_name="NAU_Doc")
        
        model = genai.GenerativeModel('gemini-3-flash-preview') # Updated to stable model name if needed, or keep 'gemini-3-flash-preview'
        
        prompt = """
        Analyze this document. Identify if it is a 'Tour Approval', 'Salary Slip', or 'Map Screenshot'.
        
        1. If **Tour Approval** (looks like "Online Tour Management System"):
           - Extract 'type': 'tour_approval'.
           - Extract 'system_no': The long number usually below a barcode or labeled "Tour ID/System No" (e.g., 21781756377236).
           - Extract 'user_details': { 'name', 'designation', 'budget_head' (B.H.) } if visible.
           - Extract 'trips': A list of journeys. For each:
             - departure_date (DD/MM/YYYY)
             - departure_time (HH:MM)
             - departure_place (City/Campus)
             - arrival_date (DD/MM/YYYY)
             - arrival_time (HH:MM)
             - arrival_place
             - mode_of_journey
             - purpose (Extract the specific reason/course name).
        
        2. If **Map Screenshot** (Google Maps):
           - Extract 'type': 'map_data'.
           - Extract 'distance_km': Numeric value of total distance (e.g., 142).
           - Extract 'travel_time': Time string (e.g., "3 hr 15 min").
           - Extract 'locations': Start and End points if visible.
        
        3. If **Salary Slip**:
           - Extract 'type': 'salary'.
           - Extract 'basic_pay'.
        
        Return ONLY valid JSON.
        """
        
        response = model.generate_content([sample_file, prompt])
        text = response.text.strip()
        if text.startswith('```json'):
            text = text.replace('```json', '').replace('```', '')
        return json.loads(text)
        
    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        return None
    finally:
        os.remove(tmp_path)

def generate_word_doc(tour_data, user_details):
    doc = Document()
    set_landscape(doc)
    
    # --- STYLES ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # --- HEADER ---
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("TOUR DIARY")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.underline = True
    
    # Determine Month Range
    dates = [t['departure_date'] for t in tour_data if t.get('departure_date')]
    month_str = ""
    if dates:
        try:
            date_objs = [datetime.strptime(d, "%d/%m/%Y") for d in dates]
            min_date = min(date_objs)
            max_date = max(date_objs)
            if min_date.month == max_date.month and min_date.year == max_date.year:
                month_str = f"Month: {min_date.strftime('%B-%Y')}"
            else:
                month_str = f"Month: {min_date.strftime('%B-%Y')} to {max_date.strftime('%B-%Y')}"
        except:
            pass

    # --- UPDATED HEADER LAYOUT (Table: 2 Rows, 3 Cols) ---
    header_table = doc.add_table(rows=2, cols=3)
    header_table.autofit = True
    
    # Row 1: Name (Left), Designation (Center), Department (Right)
    r1 = header_table.rows[0].cells
    
    # Left: Name
    r1[0].text = f"Name: {user_details.get('name', 'Vaibhav Kumar Kanubhai Chaudhari')}"
    r1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Center: Designation
    r1[1].text = f"Designation: {user_details.get('designation', 'Associate Professor')}"
    r1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Right: Dept
    r1[2].text = "Dept. of Entomology, N.M.C.A., N.A.U., Navsari"
    r1[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 2: Basic Pay (Left), B.H. (Center), Month (Right)
    r2 = header_table.rows[1].cells
    
    # Left: Basic Pay
    r2[0].text = f"Basic Salary: {user_details.get('basic_pay', 'N/A')}"
    r2[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Center: B.H.
    r2[1].text = f"B.H: {user_details.get('budget_head', '303/2092')}"
    r2[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Right: Month
    r2[2].text = month_str
    r2[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- TABLE ---
    # Columns: Dep (Place, Date, Time), Arr (Place, Date, Time), Mode, KM, Purpose
    table = doc.add_table(rows=2, cols=9)
    table.style = 'Table Grid'
    
    # Header Rows
    row0 = table.rows[0].cells
    row0[0].merge(row0[2]) # Merge first 3 for Departure
    row0[0].text = "Departure"
    row0[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    row0[3].merge(row0[5]) # Merge next 3 for Arrival
    row0[3].text = "Arrival"
    row0[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Row 1 Subheaders
    hdr_cells = table.rows[1].cells
    sub_headers = ["Place", "Date", "Time", "Place", "Date", "Time", "Mode", "KM", "Purpose of Journey"]
    for i, txt in enumerate(sub_headers):
        hdr_cells[i].text = txt
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    # Fill Data
    for trip in tour_data:
        row_cells = table.add_row().cells
        
        row_cells[0].text = str(trip.get('departure_place', 'NAU, Navsari'))
        row_cells[1].text = str(trip.get('departure_date', ''))
        row_cells[2].text = str(trip.get('departure_time', ''))
        row_cells[3].text = str(trip.get('arrival_place', ''))
        row_cells[4].text = str(trip.get('arrival_date', ''))
        row_cells[5].text = str(trip.get('arrival_time', ''))
        row_cells[6].text = str(trip.get('mode_of_journey', 'Private Vehicle'))
        row_cells[7].text = str(trip.get('distance_km', ''))
        
        # --- UPDATED PURPOSE TEXT ---
        sys_no = trip.get('system_no', '')
        purpose_desc = trip.get('purpose', '')
        purpose_text = (
            f"Subject of Tour: {purpose_desc}\n"
            f"This tour was approved by the Principal, NMCA, NAU, Navsari in Online Tour management System No. {sys_no}"
        )
        row_cells[8].text = purpose_text
        
        # Formatting cells
        for cell in row_cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell != row_cells[8] else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # --- CERTIFICATE SECTION (ADDED AS REQUESTED) ---
    p_cert_title = doc.add_paragraph()
    p_cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cert = p_cert_title.add_run("Certificate")
    run_cert.bold = True
    run_cert.font.size = Pt(11) # Matching doc font size

    p_cert_text = doc.add_paragraph("This is to certify that above said TA bill is preapred based on actual journey and actual destination with shortest routes")
    p_cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- UPDATED SIGNATURE BLOCK (Table: 1 Row, 3 Cols) ---
    # Left: User | Center: Recommended | Right: Approved
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.autofit = True
    
    # COL 1: User (Left Aligned)
    cell_user = sig_table.cell(0, 0)
    p_user = cell_user.paragraphs[0]
    p_user.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_u = p_user.add_run(
        "(V. K. Chaudhari)\n"
        "Senior Acarologist\n"
        "Dept. of Entomology\n"
        "N.M. College of Agriculture\n"
        "NAU, Navsari"
    )
    run_u.bold = True
    
    # COL 2: Recommended (Center Aligned)
    cell_rec = sig_table.cell(0, 1)
    p_rec = cell_rec.paragraphs[0]
    p_rec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_r = p_rec.add_run(
        "Recommended\n\n\n"
        "Professor and Head\n"
        "Dept. of Entomology\n"
        "N. M. College of Agriculture\n"
        "NAU, Navsari"
    )
    run_r.bold = True

    # COL 3: Approved (Right Aligned)
    cell_app = sig_table.cell(0, 2)
    p_app = cell_app.paragraphs[0]
    p_app.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_a = p_app.add_run(
        "Approved\n\n\n"
        "Principal and Dean\n"
        "N. M. College of Agriculture\n"
        "NAU, Navsari"
    )
    run_a.bold = True

    return doc

# --- MAIN APP LOGIC ---

uploaded_files = st.file_uploader("Upload Documents (PDF)", 
                                  type=['pdf'], 
                                  accept_multiple_files=True,
                                  help="Upload Tour Approvals (for dates/purpose), Map Screenshots (for distance), and Salary Slip (optional).")

if uploaded_files and st.button("Generate Word Diary"):
    if not GEMINI_API_KEY:
        st.error("Please enter your Gemini API key.")
    else:
        with st.spinner("Analyzing documents & Smart Merging..."):
            
            tour_entries = []
            map_entries = []
            user_info = {}
            
            # 1. Extract Data
            for file in uploaded_files:
                data = extract_doc_data(file, GEMINI_API_KEY)
                if data:
                    dtype = data.get('type')
                    if dtype == 'salary':
                        user_info['basic_pay'] = data.get('basic_pay')
                    
                    elif dtype == 'tour_approval':
                        if 'user_details' in data:
                            u = data['user_details']
                            if u.get('name'): user_info['name'] = u['name']
                            if u.get('designation'): user_info['designation'] = u['designation']
                            if u.get('budget_head'): user_info['budget_head'] = u['budget_head']
                        
                        trips = data.get('trips', [])
                        if isinstance(trips, list):
                            for t in trips:
                                t['system_no'] = data.get('system_no', 'Unknown')
                                tour_entries.append(t)
                        elif isinstance(trips, dict):
                             trips['system_no'] = data.get('system_no', 'Unknown')
                             tour_entries.append(trips)
                             
                    elif dtype == 'map_data':
                        map_entries.append(data)

            # 2. Smart Merge
            for i, tour in enumerate(tour_entries):
                curr_dist = str(tour.get('distance_km', '0')).strip()
                if curr_dist in ['0', '', 'None'] and map_entries:
                    m_data = map_entries[0] 
                    tour['distance_km'] = m_data.get('distance_km', '0')

            # 3. Sort by Date
            try:
                tour_entries.sort(key=lambda x: datetime.strptime(x['departure_date'], "%d/%m/%Y") if x.get('departure_date') else datetime.min)
            except:
                pass

            if tour_entries:
                doc = generate_word_doc(tour_entries, user_info)
                
                bio = io.BytesIO()
                doc.save(bio)
                
                st.success("Diary Generated Successfully!")
                st.download_button(
                    label="Download Tour Diary (.docx)",
                    data=bio.getvalue(),
                    file_name="NAU_Tour_Diary_Landscape.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("No tour data found. Please upload a valid Tour Approval PDF.")
